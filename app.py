import json
import tempfile
import uuid
from datetime import date
from html import escape
from io import BytesIO
from pathlib import Path
import shutil
import base64
import re
import zipfile
import glob
import time
import copy
import math

import pandas as pd
import streamlit as st
from docxtpl import DocxTemplate, RichText

# 修復後插入文件中的占位字串；渲染完成後由程式替換成真正表格（見 _fill_product_table_after_render）
PRODUCT_TABLE_PLACEHOLDER = "[[PRODUCT_TABLE_PLACEHOLDER]]"
DECLARATION_BULLET_PLACEHOLDER = "[[DECLARATION_BULLET_PLACEHOLDER]]"


def _fill_declaration_bullets_after_render(docx_path: Path, items: list[str]) -> bool:
    """
    自我宣告穩定方案：
    - 範本放 `{{ declaration_placeholder }}`，渲染後會出現 DECLARATION_BULLET_PLACEHOLDER
    - 程式找到該段落後，複製「該段落本身」的項目符號格式，插入 N 個 bullet 段落（每個中文內容一段）

    這樣完全不依賴 Word 內的 {% for %}，不會再遇到 unknown tag/endfor。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
    except Exception:
        return False

    if not items:
        return False

    doc = Document(str(docx_path))
    body = doc._element.body

    placeholder_idx = None
    placeholder_el = None
    for i, elt in enumerate(list(body)):
        if elt.tag != qn("w:p"):
            continue
        para = Paragraph(elt, doc)
        if DECLARATION_BULLET_PLACEHOLDER in (para.text or ""):
            placeholder_idx = i
            placeholder_el = elt
            break

    if placeholder_idx is None or placeholder_el is None:
        return False

    def _force_font(para: Paragraph) -> None:
        # 強制字型：標楷體 14pt（避免複製段落後字型回到預設）
        for run in para.runs:
            try:
                run.font.name = "標楷體"
                run.font.size = Pt(14)
            except Exception:
                pass
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
            except Exception:
                pass

    # 把占位段落「複製 N 份」，每份替換成對應文字；保留 bullet/縮排/段落樣式
    for j, text in enumerate(items):
        if j == 0:
            p = Paragraph(placeholder_el, doc)
            p.text = str(text)
            _force_font(p)
        else:
            new_p = copy.deepcopy(placeholder_el)
            body.insert(placeholder_idx + j, new_p)
            p = Paragraph(new_p, doc)
            p.text = str(text)
            _force_font(p)

    # 若原段落 text setter 生成了額外 run，仍保留 bullet 格式（pPr/numPr 在段落層）
    doc.save(str(docx_path))
    return True


def _repair_broken_product_table_template(docx_path: Path) -> None:
    """
    docxtpl 的 {%tr for %} / {%tr endfor %} 放在多欄表格列時，patch_xml 可能把標籤拆到
    `</w:tr>` 外（變成 `</w:tr>{% endfor %}`），導致 Jinja 報 unknown tag 'endfor'。

    作法：偵測仍含 {%tr 或舊版 {{p.xxx}} 迴圈的「產品表」，整表改為單一段落 `{{產品表格}}`，
    實際表格改由程式用 Subdoc 產生（見 build_products_subdoc）。
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    doc = Document(str(docx_path))
    for table in list(doc.tables):
        blob = "".join(cell.text for row in table.rows for cell in row.cells)
        if PRODUCT_TABLE_PLACEHOLDER in blob and "{%tr" not in blob:
            continue
        if "{%tr" not in blob and not ("{{p.Model}}" in blob and "產品清單" in blob):
            continue

        tbl_el = table._tbl
        parent = tbl_el.getparent()
        if parent is None:
            continue
        idx = list(parent).index(tbl_el)
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        # 使用純文字占位，渲染後再由 _fill_product_table_after_render 插入真正表格
        # （避免 docxtpl Subdoc 依賴 docxcompose）
        t.text = PRODUCT_TABLE_PLACEHOLDER
        r.append(t)
        p.append(r)
        parent.insert(idx, p)
        parent.remove(tbl_el)
        doc.save(str(docx_path))
        return


def _fill_product_table_after_render(
    docx_path: Path, product_rows: list, show_qty: bool = False, show_warranty: bool = False
) -> bool:
    """
    在 docxtpl 渲染完成後，把占位文字替換成產品表（含表頭）。
    欄位數依 show_qty / show_warranty 動態調整，與 Word 模板 {% if show_qty %} 等邏輯一致。

    重要：僅當文件中真的存在 `PRODUCT_TABLE_PLACEHOLDER`（代表模板有放 `{{產品表格}}`）
    才會插入表格；否則不做任何插入，避免「不需要表格」的模板也被硬塞產品表。
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm
        from docx.text.paragraph import Paragraph
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except Exception:
        return

    doc = Document(str(docx_path))
    body = doc._element.body

    # 僅允許：找占位字串（由模板 {{產品表格}} 渲染而來，或修復流程插入）
    insert_idx = None
    remove_paragraph_el = None
    for i, elt in enumerate(list(body)):
        if elt.tag != qn("w:p"):
            continue
        para = Paragraph(elt, doc)
        if PRODUCT_TABLE_PLACEHOLDER in (para.text or ""):
            insert_idx = i
            remove_paragraph_el = elt
            break

    # 嚴格模式：找不到占位符就直接不插表
    if insert_idx is None:
        return False

    # 欄位：固定三欄 + 選配（數量、單位、保固）
    num_cols = 3 + (2 if show_qty else 0) + (1 if show_warranty else 0)

    # add_table 會先附加到文件末端，再把表格 XML 移到目標位置
    table = doc.add_table(rows=1, cols=num_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # 關閉自動縮放，避免 Word 重新分配欄寬造成跑版
    try:
        table.autofit = False
    except Exception:
        pass

    # 表格整體寬度固定 14.56 公分
    try:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.tblW
        if tbl_w is None:
            from docx.oxml import OxmlElement

            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        # dxa 單位：1 inch = 1440，1 cm = 567 (約)
        total_dxa = int(15 * 567)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(total_dxa))
    except Exception:
        pass

    # 欄寬比例（總寬 14.56 cm）：依欄數分配
    if num_cols == 3:
        ratios = [0.15, 0.2, 0.65]
    elif num_cols == 4:
        ratios = [0.15, 0.2, 0.65, 0.07]
    elif num_cols == 5:
        # 項次、型號、設備名稱、數量、單位（僅在啟用數量時為 5 欄）
        ratios = [0.15, 0.2, 0.51, 0.07, 0.07]
    elif num_cols == 6:
        ratios = [0.15, 0.2, 0.44, 0.07, 0.07, 0.07]
    else:
        ratios = [1.0 / num_cols] * num_cols
    col_widths_cm = [15 * r for r in ratios]

    def _apply_cell_widths(row_obj):
        for idx, cell in enumerate(row_obj.cells):
            try:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.tcW
                if tc_w is None:
                    from docx.oxml import OxmlElement

                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                w_cm = col_widths_cm[idx] if idx < len(col_widths_cm) else col_widths_cm[-1]
                cell_dxa = int(w_cm * 567)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(cell_dxa))
            except Exception:
                pass

    def _apply_cell_font(cell_obj):
        for p in cell_obj.paragraphs:
            # 水平置中
            try:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
            for run in p.runs:
                run.font.name = "標楷體"
                run.font.size = Pt(12)
                # Word 中文字型需設定 eastAsia
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
                except Exception:
                    pass

    def _apply_row_layout(row_obj):
        # 固定列高 0.9 公分
        try:
            row_obj.height = Cm(0.9)
            row_obj.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        except Exception:
            pass
        # 儲存格內容垂直置中
        for cell in row_obj.cells:
            try:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            except Exception:
                pass
    tbl_el = table._tbl
    body.remove(tbl_el)
    body.insert(insert_idx, tbl_el)
    if remove_paragraph_el is not None:
        body.remove(remove_paragraph_el)

    hdr = table.rows[0].cells
    hi = 0
    hdr[hi].text = "項次"
    hi += 1
    hdr[hi].text = "ACTi型號"
    hi += 1
    hdr[hi].text = "設備名稱"
    hi += 1
    if show_qty:
        hdr[hi].text = "數量"
        hi += 1
        hdr[hi].text = "單位"
        hi += 1
    if show_warranty:
        hdr[hi].text = "保固"
        hi += 1
    _apply_row_layout(table.rows[0])
    _apply_cell_widths(table.rows[0])
    for c in table.rows[0].cells:
        _apply_cell_font(c)

    for row in product_rows:
        r = table.add_row()
        ci = 0
        r.cells[ci].text = str(row.get("項次", ""))
        ci += 1
        r.cells[ci].text = str(row.get("Model", ""))
        ci += 1
        r.cells[ci].text = str(row.get("設備名稱", row.get("ProductName", "")))
        ci += 1
        if show_qty:
            r.cells[ci].text = str(row.get("數量", "") or "")
            ci += 1
            r.cells[ci].text = str(row.get("單位", "") or "")
            ci += 1
        if show_warranty:
            r.cells[ci].text = str(row.get("保固", "") or "")
            ci += 1
        _apply_row_layout(r)
        _apply_cell_widths(r)
        for c in r.cells:
            _apply_cell_font(c)
    doc.save(str(docx_path))
    return True


def _force_product_table_cell_widths(
    docx_path: Path, show_qty: bool, show_warranty: bool
) -> None:
    """
    docxtpl 無法穩定控制欄寬時，於渲染完成後使用 python-docx 強制設定「產品明細」表格欄寬。

    規格（總寬 14.55 cm）：
    1) show_qty=True 且 show_warranty=True  -> 6 欄
       [1.3, 2.48, 7.46, 1.3, 1.3, 1.3]
    2) show_qty=True 且 show_warranty=False -> 5 欄
       [1.3, 2.48, 8.59, 1.3, 1.3]
    3) 兩者皆 False                          -> 3 欄
       [1.3, 2.48, 10.85]

    注意：此函式會依「每列 row、每個 cell」逐一賦值 cell.width，避免 Word 自動擠壓。
    """
    try:
        from docx import Document
        from docx.shared import Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception:
        return

    try:
        doc = Document(str(docx_path))
    except Exception:
        return

    # 優先找「表頭包含項次/ACTi型號/設備名稱」的產品表；找不到就取第一個有足夠欄位的表格。
    target_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = table.rows[0].cells
        header_texts = [(c.text or "").strip() for c in header_cells]
        if len(header_texts) >= 3:
            if ("項次" in header_texts[0]) and ("ACTi" in header_texts[1]) and ("設備" in header_texts[2]):
                target_table = table
                break
    if target_table is None:
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                texts = [(c.text or "").strip() for c in table.rows[0].cells]
                if texts and "項次" in texts[0]:
                    target_table = table
                    break
    if target_table is None:
        return

    # 依勾選狀態與「實際欄數」決定欄寬（總寬 14.55cm）
    actual_cols = 0
    try:
        actual_cols = len(target_table.rows[0].cells) if target_table.rows else 0
    except Exception:
        actual_cols = 0

    # 依「實際欄數」強制覆寫（你的新規格只定義 3/5/6 欄）
    # 3 欄：項次、型號、名稱
    # 5 欄：項次、型號、名稱、數量、單位
    # 6 欄：項次、型號、名稱、數量、單位、保固
    if actual_cols == 6:
        widths_cm = [1.3, 2.5, 7.3, 1.3, 1.3, 1.3]
    elif actual_cols == 5:
        widths_cm = [1.3, 2.5, 8.6, 1.3, 1.3]
    elif actual_cols == 3:
        widths_cm = [1.5, 2.5, 11.2]
    else:
        return

    expected_cols = len(widths_cm)

    # 關閉 Word 自動調整（並盡可能把 XML 設成 fixed layout）
    try:
        target_table.autofit = False
    except Exception:
        pass
    try:
        tbl_pr = target_table._tbl.tblPr
        # <w:tblLayout w:type="fixed"/>
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")
    except Exception:
        pass

    # 固定整表寬度（依欄寬加總，約 15cm 左右；dxa：1 cm = 567）
    try:
        tbl_pr = target_table._tbl.tblPr
        tbl_w = tbl_pr.tblW
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        total_cm = float(sum(widths_cm))
        total_dxa = int(total_cm * 567)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(total_dxa))
    except Exception:
        pass

    # 同步設定欄網格 <w:tblGrid>，降低 Word 重新分配欄寬的機率
    try:
        tbl = target_table._tbl
        grid = tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            tbl.insert(0, grid)
        # 先清空舊的 gridCol
        for child in list(grid):
            try:
                grid.remove(child)
            except Exception:
                pass
        for w_cm in widths_cm:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(float(w_cm) * 567)))
            grid.append(gc)
    except Exception:
        pass

    # 必須同時設定每個 column 的 width
    try:
        for i, col in enumerate(target_table.columns):
            if i >= expected_cols:
                break
            try:
                col.width = Cm(float(widths_cm[i]))
            except Exception:
                pass
    except Exception:
        pass

    # 依規格逐列逐 cell 強制設定欄寬
    for row in target_table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= expected_cols:
                break
            w_cm = widths_cm[col_idx]
            try:
                cell.width = Cm(w_cm)
            except Exception:
                pass
            # 也同步寫入 tcW（讓 Word XML 層級也一致）
            try:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.tcW
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(w_cm * 567)))
            except Exception:
                pass

    try:
        doc.save(str(docx_path))
    except Exception:
        return


def _taiwan_date_strings(d: date) -> dict:
    """將日期轉成台灣常用格式(西元 + 民國)並提供多種鍵名給模板使用。"""
    year = d.year
    month = d.month
    day = d.day
    roc_year = year - 1911

    # yyMMDD：西元兩位年份 + 月 + 日
    yy = year % 100

    # 民國的 yyMMDD：民國年份取最後兩位
    roc_yy = roc_year % 100

    # 常見台灣寫法：YYYY/MM/DD 或 民國年/YYYY/MM/DD
    date_greg_dashless = f"{year:04d}{month:02d}{day:02d}"  # 例如 20260320
    date_greg_slash = f"{year:04d}/{month:02d}/{day:02d}"  # 例如 2026/03/20
    date_roc_slash = f"{roc_year:03d}/{month:02d}/{day:02d}"  # 例如 115/03/20
    date_yyMMdd = f"{yy:02d}{month:02d}{day:02d}"  # 例如 260320
    date_roc_yyMMdd = f"{roc_yy:02d}{month:02d}{day:02d}"  # 例如 150320

    # 中文常見寫法：民國115年03月20日
    date_greg_roc_cn = f"{roc_year}年{month:02d}月{day:02d}日"

    return {
        "date_greg_dashless": date_greg_dashless,
        "date_greg_slash": date_greg_slash,
        "date_roc_slash": date_roc_slash,
        "date_yyMMdd": date_yyMMdd,
        "date_roc_yyMMdd": date_roc_yyMMdd,
        # 民國年/月/日(中文格式)
        "date_roc_year_month_day_cn": date_greg_roc_cn,
        # 也提供中文鍵名(若你的 Word 模板用中文變數)
        "日期_西元_YYYYMMDD": date_greg_dashless,
        "日期_西元_YYYY/MM/DD": date_greg_slash,
        "日期_民國_YYYY/MM/DD": date_roc_slash,
        "日期_西元_yyMMDD": date_yyMMdd,
        "日期_民國_yyMMDD": date_roc_yyMMdd,
        "日期_民國_年月日": date_greg_roc_cn,
    }


def _normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    嘗試把 Excel 的欄位名稱標準化成我們需要的欄位：
    - Model
    - ProductName (設備名稱/產品名稱)
    - 單位（選用，若無則補空字串）
    - 保固（選用，若無則補空字串）
    """
    # 先做一個「去空白+小寫」的索引，方便模糊比對
    col_map = {}
    for c in df.columns:
        key = str(c).strip().lower()
        col_map[key] = c

    def pick(*candidates: str):
        for cand in candidates:
            cand_key = cand.strip().lower()
            if cand_key in col_map:
                return col_map[cand_key]
        return None

    model_col = pick("model", "機型", "型號", "型式")
    product_col = pick(
        "productname",
        "產品名稱",
        "設備名稱",
        "品名",
        "設備名",
    )
    # 選配欄位：名稱可能因地區或表頭習慣略有不同
    unit_col = pick("單位", "unit", "計量單位", "计量单位")
    warranty_col = pick("保固", "warranty", "保固期", "保固資訊", "保固说明")

    if model_col is None or product_col is None:
        raise ValueError(
            "產品資料庫欄位不足，找不到所需欄位。請確認 Excel 至少包含 Model 與 ProductName(設備/產品名稱) 欄位。\n"
            f"目前欄位：{list(df.columns)}"
        )

    rename_map = {model_col: "Model", product_col: "ProductName"}
    if unit_col is not None:
        rename_map[unit_col] = "單位"
    if warranty_col is not None:
        rename_map[warranty_col] = "保固"

    out = df.rename(columns=rename_map)
    if "單位" not in out.columns:
        out["單位"] = ""
    if "保固" not in out.columns:
        out["保固"] = ""

    return out[["Model", "ProductName", "單位", "保固"]]


@st.cache_data(show_spinner=False)
def load_product_db(product_db_path: str) -> pd.DataFrame:
    """從 `產品資料庫.xlsx` 載入產品資料，並做欄位標準化。"""
    df = pd.read_excel(product_db_path, engine="openpyxl")
    df = _normalize_columns(df)

    # 清理空白與缺值，避免 selectbox 顯示奇怪的值
    df["Model"] = df["Model"].astype(str).str.strip()
    df["ProductName"] = df["ProductName"].astype(str).str.strip()
    df["單位"] = df["單位"].fillna("").astype(str).str.strip()
    df["保固"] = df["保固"].fillna("").astype(str).str.strip()
    df = df[(df["Model"] != "") & (df["ProductName"] != "")]

    return df


@st.cache_data(show_spinner=False)
def list_docx_templates(base_dir: str) -> dict:
    """
    掃描模板檔案，找出 `.docx` 模板並回傳：
      { "模板名(檔名不含副檔名)": "模板檔路徑" }
    預設會掃描：
      1) 專案根目錄
      2) `templates/` 子資料夾(若存在)
    """
    base = Path(base_dir)

    templates: dict[str, str] = {}

    # 以遞迴方式掃描，確保模板可能放在子資料夾(例如 `ACTi_Doc_Generator/`)時也能找到。
    for p in sorted(base.rglob("*.docx")):
        # 避免誤拾取臨時檔(例如 Word 會產生 ~$xxx.docx)
        if p.name.startswith("~$"):
            continue
        templates[p.stem] = str(p)

    return templates


def sanitize_filename(s: str) -> str:
    """把檔名中的非法字元替換成底線，避免 Windows 存檔失敗。"""
    s = s.strip()
    for ch in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
        s = s.replace(ch, "_")
    # 避免檔名過長
    return s[:120] if len(s) > 120 else s


def build_context(
    use_unit: str,
    vendor_name: str,
    project_name: str,
    distributor: str,
    other_info: str,
    d: date,
    products: list,
    show_qty: bool,
    show_warranty: bool,
) -> dict:
    """
    組合給 docxtpl 使用的 context 字典。
    這裡同時提供中文與英文 key，讓你可以依模板實際用的變數名稱選擇對應。
    show_qty / show_warranty：供 Word 模板 {% if show_qty %} 決定是否顯示數量／單位／保固欄。
    """
    use_unit = (use_unit or "").strip()
    vendor_name = (vendor_name or "").strip()
    project_name = (project_name or "").strip()
    distributor = (distributor or "").strip()
    other_info = (other_info or "").strip()

    # 舊模板相容：以「案件名稱」對應舊的案名；其餘不再使用多行補充
    case_all = project_name
    case_title = project_name
    case_extra = ""

    date_dict = _taiwan_date_strings(d)

    # 產品清單(給模板 loop 用)
    # 你可以在 Word 模板中用 {{ for p in products }} 或 {{產品清單}} 的方式對應。
    product_rows = []
    for i, item in enumerate(products, start=1):
        # 選配欄位：未啟用時存空字串，模板與表格後處理仍可安全讀取
        qty_val = item.get("數量", "")
        unit_val = item.get("單位", "")
        war_val = item.get("保固", "")
        product_rows.append(
            {
                "項次": i,
                "Model": item["Model"],
                "設備名稱": item["ProductName"],
                "數量": qty_val if qty_val is not None else "",
                "單位": unit_val if unit_val is not None else "",
                "保固": war_val if war_val is not None else "",
                # 英文 key(也提供給可能使用英文變數的模板)
                "index": i,
                "model": item["Model"],
                "product_name": item["ProductName"],
                "qty": qty_val if qty_val is not None else "",
                "unit": unit_val if unit_val is not None else "",
                "warranty": war_val if war_val is not None else "",
            }
        )

    context = {
        # 基本資訊 - 中文 key
        "廠商名稱": vendor_name,
        "使用單位": use_unit,
        "案名": case_title,
        # 提供「整段內容」給模板直接用，避免只顯示第一行造成看起來像沒填到
        "案名全文": case_all,
        "案名_全部": case_all,
        "補充資訊": case_extra,
        **date_dict,
        "產品清單": product_rows,
        # 若模板有 {{產品表格}}，先渲染成占位字串，稍後由程式替換成真正 Word 表格
        "產品表格": PRODUCT_TABLE_PLACEHOLDER,
        # 額外別名：讓模板不管用哪個 key 都取得到同一份產品清單
        "model": product_rows,
        "model_table": product_rows,
        # 基本資訊 - 英文 key
        "use_unit": use_unit,
        "vendor_name": vendor_name,
        "project_name": project_name,
        "distributor": distributor,
        "other_info": other_info,
        # 也提供一份中文 key（新模板也可直接用）
        "案件名稱": project_name,
        "經銷商名稱": distributor,
        "其他資訊": other_info,
        # 舊 key 兼容（若你舊模板仍在用）
        "unit_name": use_unit,
        "case_title": case_title,
        "case_text_all": case_all,
        "case_extra": case_extra,
        "products": product_rows,
        "date_greg_slash": date_dict["date_greg_slash"],
        "date_roc_slash": date_dict["date_roc_slash"],
        "date_greg_dashless": date_dict["date_greg_dashless"],
        "date_roc_year_month_day_cn": date_dict["date_roc_year_month_day_cn"],
        # 布林：Word 模板可用 {% if show_qty %} / {% if show_warranty %}
        "show_qty": bool(show_qty),
        "show_warranty": bool(show_warranty),
    }

    return context


def render_doc(template_path: str, context: dict) -> bytes:
    """使用 docxtpl 渲染模板，並回傳生成後 docx 的 bytes。"""
    # Windows 上若模板檔在 Word 中被開啟，可能會被鎖住導致 docxtpl/zip 無法讀取。
    # 因此這裡先複製一份到 temp，確保讀取目標不會跟外部程式互相搶鎖。
    tpl_tmp_path = Path(tempfile.gettempdir()) / f"tpl_{uuid.uuid4().hex}.docx"
    try:
        shutil.copyfile(template_path, tpl_tmp_path)
        # 自動修復含 {%tr %} 的表格（避免 docxtpl patch_xml 產生孤立的 {% endfor %}）
        _repair_broken_product_table_template(tpl_tmp_path)
        tpl = DocxTemplate(str(tpl_tmp_path))
        tpl.render(context)
    finally:
        try:
            tpl_tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

    tmp_dir = Path(tempfile.gettempdir())
    out_name = f"generated_{uuid.uuid4().hex}.docx"
    out_path = tmp_dir / out_name
    tpl.save(str(out_path))
    # 將占位字串替換成實際表格（欄位數依 context 的 show_qty / show_warranty）
    rows = context.get("產品清單") or context.get("products") or []
    inserted = _fill_product_table_after_render(
        out_path,
        rows,
        bool(context.get("show_qty", False)),
        bool(context.get("show_warranty", False)),
    )
    # 只有真的插入了產品表，才做欄寬強制，避免誤調整其他表格
    if inserted:
        _force_product_table_cell_widths(
            out_path,
            bool(context.get("show_qty", False)),
            bool(context.get("show_warranty", False)),
        )

    # 自我宣告：若有提供 declaration_items 且範本內存在 declaration_placeholder，則插入 bullet 段落
    try:
        decl_items = context.get("declaration_items") or []
        if isinstance(decl_items, list) and decl_items:
            _fill_declaration_bullets_after_render(out_path, [str(x) for x in decl_items if str(x).strip()])
    except Exception:
        pass

    data = out_path.read_bytes()
    try:
        out_path.unlink(missing_ok=True)
    except Exception:
        # 若刪除失敗不影響使用
        pass
    return data


# 缺漏檢查規則（檔名或宣告名稱關鍵字 → 必填維度）
# party: 廠商名稱或使用單位擇一；cart: 購物車內至少一筆產品
VALIDATION_RULES: dict[str, tuple[str, ...]] = {
    "停產": ("cart",),
    "出廠證明": ("party", "cart"),
    "__default_template__": ("party", "cart"),
    "__default_declaration__": ("party",),
}


def _missing_fields_for_document(
    name_for_rules: str,
    *,
    use_unit: str,
    vendor_name: str,
    cart: list,
    is_declaration_row: bool,
) -> list[str]:
    """
    依 VALIDATION_RULES 與檔名／宣告名稱關鍵字判斷缺漏欄位（與 ensure_rendered 邏輯對齊）。
    - 「停產」：須購物車內有產品。
    - 「出廠證明」：須（廠商名稱或使用單位）且購物車有產品。
    - 自我宣告（一般）：須廠商／使用單位擇一。
    - 其他 Word 範本：須廠商／使用單位擇一，且購物車有產品。
    """
    s = name_for_rules or ""
    has_party = bool(str(use_unit).strip() or str(vendor_name).strip())
    has_cart = bool(cart and len(cart) > 0)

    if "停產" in s:
        missing: list[str] = []
        if not has_cart:
            missing.append("購物車內產品（Model）")
        return missing

    if "出廠證明" in s:
        missing = []
        if not has_party:
            missing.append("廠商名稱或使用單位（擇一填寫）")
        if not has_cart:
            missing.append("購物車內產品（Model）")
        return missing

    if is_declaration_row:
        if not has_party:
            return ["廠商名稱或使用單位（擇一填寫）"]
        return []

    missing = []
    if not has_party:
        missing.append("廠商名稱或使用單位（擇一填寫）")
    if not has_cart:
        missing.append("購物車內產品（Model）")
    return missing


def collect_validation_issues(
    *,
    selected_template_paths: list[str],
    selected_declarations: list[str],
    use_unit: str,
    vendor_name: str,
    cart: list,
) -> list[tuple[str, list[str]]]:
    """
    回傳 [(顯示名稱, 缺漏欄位列表), ...] 僅含未通過驗證的文件。
    """
    issues: list[tuple[str, list[str]]] = []
    for decl in selected_declarations:
        label = f"{decl}（自我宣告）"
        miss = _missing_fields_for_document(
            decl,
            use_unit=use_unit,
            vendor_name=vendor_name,
            cart=cart,
            is_declaration_row=True,
        )
        if miss:
            issues.append((label, miss))
    for path in selected_template_paths:
        stem = Path(path).stem
        miss = _missing_fields_for_document(
            stem,
            use_unit=use_unit,
            vendor_name=vendor_name,
            cart=cart,
            is_declaration_row=False,
        )
        if miss:
            issues.append((stem, miss))
    return issues


def render_missing_info_tracker(
    *,
    selected_template_paths: list[str],
    selected_declarations: list[str],
    use_unit: str,
    vendor_name: str,
    cart: list,
) -> None:
    """缺漏資料提示區塊（Missing Info Tracker）。"""
    st.subheader("缺漏資料提示")
    n_sel = len(selected_template_paths) + len(selected_declarations)
    if n_sel == 0:
        st.info("請於「區塊二：文件類型選擇」勾選至少一份要產生的文件。")
        return

    issues = collect_validation_issues(
        selected_template_paths=selected_template_paths,
        selected_declarations=selected_declarations,
        use_unit=use_unit,
        vendor_name=vendor_name,
        cart=cart,
    )
    if not issues:
        st.success("所有資料已就緒，可進行預覽與下載！")
        return

    for doc_label, missing in issues:
        miss_txt = "、".join(missing)
        st.warning(f"要生成【{doc_label}】，您還缺少：{miss_txt}。")


def _empty_cart_preview_dataframe(show_qty: bool, show_warranty: bool) -> pd.DataFrame:
    """
    購物車尚無資料時，仍依勾選顯示完整欄位標題（數量／單位／保固），方便使用者預覽表結構。
    """
    cols = ["項次", "Model", "設備名稱"]
    if show_qty:
        cols.extend(["數量", "單位"])
    if show_warranty:
        cols.append("保固")
    return pd.DataFrame(columns=cols)


def _cart_to_editor_dataframe(cart: list, show_qty: bool, show_warranty: bool) -> pd.DataFrame:
    """將 session 購物車轉成 data_editor 用 DataFrame（數量預設為 1 供右側表格編輯）。"""
    rows = []
    for i, p in enumerate(cart):
        row = {
            "項次": i + 1,
            "Model": p.get("Model", ""),
            "設備名稱": p.get("ProductName", ""),
        }
        if show_qty:
            q = p.get("數量", "")
            if q == "" or q is None:
                qv = 1
            else:
                try:
                    qv = int(float(q))
                except (TypeError, ValueError):
                    qv = 1
            row["數量"] = qv
            row["單位"] = str(p.get("單位", "") or "")
        if show_warranty:
            row["保固"] = str(p.get("保固", "") or "")
        rows.append(row)
    return pd.DataFrame(rows)


def _sync_qty_from_editor(edited_df: pd.DataFrame, cart: list, show_qty: bool) -> None:
    """將「目前已加入的產品」表格中編輯的數量寫回購物車；單位／保固仍來自 Excel／加入時帶入。"""
    if not show_qty or edited_df.empty or not cart:
        return
    if "數量" not in edited_df.columns:
        return
    n = min(len(edited_df), len(cart))
    for i in range(n):
        try:
            q = edited_df.iloc[i]["數量"]
            cart[i]["數量"] = max(1, int(float(q)))
        except (TypeError, ValueError, KeyError):
            cart[i]["數量"] = 1


def docx_bytes_to_html_preview(docx_bytes: bytes) -> str | None:
    """
    將 docx bytes 轉成 HTML 供 Streamlit 預覽。
    需要額外套件 `mammoth`。
    若未安裝/轉換失敗，回傳 None 讓 UI 降級。
    """
    try:
        import mammoth  # 不是內建套件，需要使用者自行安裝
    except Exception:
        return None

    try:
        res = mammoth.convert_to_html(BytesIO(docx_bytes))
        return res.value
    except Exception:
        return None


def docx_bytes_to_pdf_preview_base64(docx_bytes: bytes) -> str | None:
    """
    將 docx bytes 轉成 PDF，並回傳 base64 字串供 iframe 預覽。
    優先使用 `docx2pdf`（Windows 可透過 Word 進行轉檔）。

    若沒有安裝 `docx2pdf` 或轉檔失敗，回傳 None 讓 UI 自動降級。
    """
    pdf_bytes, _err = docx_bytes_to_pdf_bytes(docx_bytes)
    if not pdf_bytes:
        return None
    return base64.b64encode(pdf_bytes).decode("utf-8")


def init_session_state():
    """初始化 session_state，確保購物車與表單鍵值存在，避免 KeyError。"""
    if "cart_products" not in st.session_state:
        # 每個元素：Model, ProductName，以及選配的 數量、單位、保固
        st.session_state.cart_products = []
    # 供頁首進度條在「同一輪、widget 尚未建立前」讀取上一輪輸入值
    # 區塊一：基本資訊（新版欄位）
    if "inp_use_unit" not in st.session_state:
        st.session_state.inp_use_unit = ""
    if "inp_vendor_name" not in st.session_state:
        st.session_state.inp_vendor_name = ""
    if "inp_project_name" not in st.session_state:
        st.session_state.inp_project_name = ""
    if "inp_distributor" not in st.session_state:
        st.session_state.inp_distributor = ""
    if "inp_other_info" not in st.session_state:
        st.session_state.inp_other_info = ""

    # 舊欄位（保留避免既有 session_state 造成錯誤；新 UI 不再使用）
    if "inp_vendor" not in st.session_state:
        st.session_state.inp_vendor = ""
    if "inp_doc_types" not in st.session_state:
        st.session_state.inp_doc_types = []
    if "selected_template_paths" not in st.session_state:
        st.session_state.selected_template_paths = []
    if "selected_declarations" not in st.session_state:
        st.session_state.selected_declarations = []
    if "last_pdf_auto_signature" not in st.session_state:
        st.session_state.last_pdf_auto_signature = ""
    if "last_pdf_auto_b64" not in st.session_state:
        st.session_state.last_pdf_auto_b64 = None
    if "last_pdf_auto_attempt_sig" not in st.session_state:
        st.session_state.last_pdf_auto_attempt_sig = ""
    if "last_pdf_auto_attempt_ts" not in st.session_state:
        st.session_state.last_pdf_auto_attempt_ts = 0.0
    if "last_auto_signature" not in st.session_state:
        st.session_state.last_auto_signature = ""
    if "last_auto_results" not in st.session_state:
        st.session_state.last_auto_results = {}


def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> tuple[bytes | None, str | None]:
    """
    將 docx 轉成 PDF bytes（docx2pdf，通常需本機安裝 Microsoft Word）。
    回傳 (pdf_bytes, 錯誤說明)；成功時錯誤說明為 None。
    """
    try:
        import docx2pdf
    except Exception:
        return None, "未安裝 docx2pdf，請執行：pip install docx2pdf"

    tmp_dir = Path(tempfile.gettempdir())
    uid = uuid.uuid4().hex
    in_docx = tmp_dir / f"dl_{uid}.docx"
    out_pdf = tmp_dir / f"dl_{uid}.pdf"

    try:
        in_docx.write_bytes(docx_bytes)
        try:
            docx2pdf.convert(str(in_docx), str(out_pdf))
        except Exception:
            docx2pdf.convert(str(in_docx), str(tmp_dir))
            candidates = sorted(tmp_dir.glob(f"dl_{uid}*.pdf"))
            if candidates:
                out_pdf = candidates[0]

        if not out_pdf.exists():
            return None, "Word 轉 PDF 未產生檔案（請確認已安裝 Microsoft Word 且 docx2pdf 可正常運作）。"
        return out_pdf.read_bytes(), None
    except Exception as e:
        return None, f"Word 轉 PDF 失敗（通常需安裝 Microsoft Word）：{e}"
    finally:
        try:
            in_docx.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            out_pdf.unlink(missing_ok=True)
        except Exception:
            pass


def pdf_bytes_to_png_bytes(pdf_bytes: bytes, *, dpi: float = 144.0) -> tuple[bytes | None, str | None]:
    """
    將 PDF 第一頁轉成 PNG（PyMuPDF / fitz）。
    多頁文件僅輸出第一頁，避免檔案過大。
    """
    try:
        import fitz  # PyMuPDF
    except Exception:
        return None, "未安裝 PyMuPDF，請執行：pip install pymupdf"

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count < 1:
            doc.close()
            return None, "PDF 無任何頁面。"
        page = doc.load_page(0)
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png_bytes = pix.tobytes("png")
        doc.close()
        return png_bytes, None
    except Exception as e:
        return None, f"PDF 轉 PNG 失敗：{e}"


def result_key_for_preview_pick(
    preview_pick: str,
    decl_label: str | None,
    decl_selected: list[str],
    selected_paths: list[str],
) -> str | None:
    """對應「選擇預覽文件」選項與 rendered_results 的 dict key。"""
    if not preview_pick or not preview_pick.strip():
        return None
    if decl_label and preview_pick == decl_label:
        if not decl_selected:
            return None
        return f"__DECL__::{','.join(decl_selected)}"
    for p in selected_paths:
        if Path(p).stem == preview_pick:
            return p
    return None


def ensure_rendered_docx_results(
    *,
    use_unit: str,
    vendor_name: str,
    project_name: str,
    distributor: str,
    other_info: str,
    chosen_date: date,
    assets_dir: Path,
    base_dir: Path,
    declaration_df,
) -> tuple[dict[str, dict], str | None]:
    """
    依目前 session 勾選與購物車，產生或讀取快取的 docx 結果。
    回傳 (rendered_results, render_error)。
    """
    has_party = bool((use_unit or "").strip() or (vendor_name or "").strip())
    has_cart = bool(st.session_state.cart_products)
    decl_selected = st.session_state.get("selected_declarations", []) or []
    has_any_doc = bool(st.session_state.get("selected_template_paths", []) or []) or bool(decl_selected)
    can_download = has_party and has_any_doc and (has_cart or bool(decl_selected))

    if not can_download:
        return {}, None

    show_qty_flag = bool(st.session_state.get("opt_qty", False))
    show_warranty_flag = bool(st.session_state.get("opt_warranty", False))
    context = build_context(
        use_unit=use_unit,
        vendor_name=vendor_name,
        project_name=project_name,
        distributor=distributor,
        other_info=other_info,
        d=chosen_date,
        products=st.session_state.cart_products if has_cart else [],
        show_qty=show_qty_flag,
        show_warranty=show_warranty_flag,
    )
    if decl_selected and declaration_df is not None and (not declaration_df.empty):
        parts = declaration_df[declaration_df["宣告"].isin(decl_selected)]["中文內容"].tolist()
        parts = [str(x).strip() for x in parts if str(x).strip()]
        if parts:
            context["declaration_items"] = parts
            context["declaration_placeholder"] = DECLARATION_BULLET_PLACEHOLDER

    signature_payload = {
        "use_unit": use_unit,
        "vendor_name": vendor_name,
        "project_name": project_name,
        "distributor": distributor,
        "other_info": other_info,
        "decl_selected": decl_selected,
        "date": str(chosen_date),
        "paths": sorted(st.session_state.get("selected_template_paths", []) or []),
        "cart": st.session_state.cart_products if has_cart else [],
        "show_qty": show_qty_flag,
        "show_warranty": show_warranty_flag,
    }
    signature = json.dumps(signature_payload, ensure_ascii=False, sort_keys=True)

    if st.session_state.get("last_auto_signature") != signature:
        st.session_state.last_auto_signature = signature
        st.session_state.last_auto_results = {}
        rendered_results: dict[str, dict] = {}
        try:
            if has_cart:
                for path in st.session_state.get("selected_template_paths", []) or []:
                    out = render_doc(path, context)
                    safe_vendor = sanitize_filename(vendor_name) if vendor_name else "廠商"
                    safe_date = sanitize_filename(_taiwan_date_strings(chosen_date)["date_yyMMdd"])
                    file_name = f"{sanitize_filename(Path(path).stem)}_{safe_vendor}_{safe_date}.docx"
                    rendered_results[path] = {"bytes": out, "filename": file_name}

            if decl_selected and context.get("declaration_items"):
                decl_tpl = None
                for root in [assets_dir, base_dir]:
                    p = root / "規格符合證明書.docx"
                    if p.exists():
                        decl_tpl = str(p)
                        break
                if decl_tpl:
                    out = render_doc(decl_tpl, context)
                    name_prefix = ",".join(decl_selected)
                    decl_filename = f"{name_prefix}自我宣告.docx"
                    rendered_results[f"__DECL__::{name_prefix}"] = {"bytes": out, "filename": decl_filename}

            st.session_state.last_auto_results = rendered_results
        except Exception as e:
            return {}, str(e)

    return dict(st.session_state.get("last_auto_results", {})), None


def scan_and_classify_docx(base_dir: Path) -> dict[str, dict[str, str]]:
    """
    動態掃描 app.py 同層與 templates 目錄下的 docx，並依檔名分類：
    - 自我宣告：含「自我宣告」
    - 證明文件：含「證明書」
    - 英文文件：檔名(不含副檔名)僅 ASCII
    - 其他文件：其餘
    """
    categories = {"自我宣告": {}, "證明文件": {}, "英文文件": {}, "其他文件": {}}
    # 你把檔案放在下一層資料夾「資料庫&文件範本」：優先掃那裡，並保留原本目錄的 fallback
    tpl_root = base_dir / "資料庫&文件範本"
    patterns = [
        str(tpl_root / "*.docx"),
        str(tpl_root / "templates" / "*.docx"),
        str(base_dir / "*.docx"),
        str(base_dir / "templates" / "*.docx"),
    ]
    all_paths: set[Path] = set()
    for pattern in patterns:
        for fp in glob.glob(pattern):
            p = Path(fp)
            if p.name.startswith("~$"):
                continue
            all_paths.add(p)

    for p in sorted(all_paths, key=lambda x: x.stem.lower()):
        stem = p.stem
        if "自我宣告" in stem:
            categories["自我宣告"][stem] = str(p)
        elif "證明書" in stem:
            categories["證明文件"][stem] = str(p)
        elif re.match(r"^[\x00-\x7F]+$", stem):
            categories["英文文件"][stem] = str(p)
        else:
            categories["其他文件"][stem] = str(p)
    return categories


@st.cache_data(show_spinner=False)
def load_declarations_xlsx(xlsx_path: str) -> pd.DataFrame:
    """
    讀取 `自我宣告.xlsx`，需包含：
    - 宣告
    - 中文內容
    會排除空值並回傳整理後的 DataFrame。
    """
    df = pd.read_excel(xlsx_path, engine="openpyxl")
    # 欄位防呆：去空白
    df.columns = [str(c).strip() for c in df.columns]
    if ("宣告" not in df.columns) or ("中文內容" not in df.columns):
        raise ValueError(f"自我宣告.xlsx 欄位不足，需包含『宣告』『中文內容』；目前欄位：{list(df.columns)}")
    out = df[["宣告", "中文內容"]].copy()
    out["宣告"] = out["宣告"].fillna("").astype(str).str.strip()
    out["中文內容"] = out["中文內容"].fillna("").astype(str).str.strip()
    out = out[(out["宣告"] != "") & (out["中文內容"] != "")]
    # 去重：以宣告為 key，保留第一筆
    out = out.drop_duplicates(subset=["宣告"], keep="first")
    return out


def build_live_preview_html(vendor_name: str, case_text: str, chosen_date: date, cart_df: pd.DataFrame) -> str:
    """右側即時預覽：A4 白紙風格 + 基本資訊 + 動態產品表格。"""
    d = _taiwan_date_strings(chosen_date)
    case_lines = [ln.strip() for ln in (case_text or "").splitlines() if ln.strip()]
    case_title = case_lines[0] if case_lines else ""
    case_extra = "<br>".join(escape(x) for x in case_lines[1:]) if len(case_lines) > 1 else ""
    if cart_df.empty:
        table_html = "<p style='color:#64748b;'>尚無產品資料</p>"
    else:
        table_html = cart_df.to_html(index=False, escape=True, classes="preview-table")
    return f"""
<style>
.a4-wrap {{
  background:#ffffff;
  border:1px solid #e5e7eb;
  box-shadow:0 8px 24px rgba(15,23,42,.08);
  border-radius:10px;
  padding:28px 30px;
  min-height:760px;
}}
.a4-title {{ font-size:1.2rem; font-weight:700; margin-bottom:12px; }}
.a4-meta {{ color:#334155; line-height:1.7; margin-bottom:16px; }}
.preview-table {{ width:100%; border-collapse:collapse; font-size:.88rem; }}
.preview-table th, .preview-table td {{
  border:1px solid #cbd5e1; padding:6px 8px; text-align:center; vertical-align:middle;
}}
.preview-table th {{ background:#f8fafc; }}
</style>
<div class="a4-wrap">
  <div class="a4-title">文件預覽</div>
  <div class="a4-meta">
    <div><b>廠商名稱：</b>{escape(vendor_name) if vendor_name else "-"}</div>
    <div><b>案名：</b>{escape(case_title) if case_title else "-"}</div>
    <div><b>補充：</b>{case_extra if case_extra else "-"}</div>
    <div><b>日期：</b>西元 {escape(d["date_greg_slash"])} / 民國 {escape(d["date_roc_slash"])}</div>
  </div>
  <div>{table_html}</div>
</div>
"""


def main():
    st.set_page_config(page_title="ACTi 文件自動生成中心", layout="wide")
    init_session_state()
    base_dir = Path(__file__).resolve().parent
    assets_dir = base_dir / "資料庫&文件範本"
    product_db_candidates = ["產品資料庫.xlsx", "model_list.xlsx"]

    st.title("ACTi 文件自動生成中心")

    # Excel 路徑防呆（優先：資料庫&文件範本/；fallback：原目錄）
    product_db_path = None
    for root in [assets_dir, base_dir]:
        for name in product_db_candidates:
            p = root / name
            if p.exists():
                product_db_path = p
                break
        if product_db_path is not None:
            break
    if product_db_path is None:
        fuzzy = []
        for root in [assets_dir, base_dir]:
            fuzzy.extend(sorted(root.glob("*產品資料庫*.xlsx*")))
            fuzzy.extend(sorted(root.glob("*model_list*.xlsx*")))
        fuzzy = [p for p in fuzzy if p.name and not p.name.startswith("~$")]
        for p in fuzzy:
            if "產品資料庫" in p.name:
                product_db_path = p
                break
        if product_db_path is None and fuzzy:
            product_db_path = fuzzy[0]

    categorized = scan_and_classify_docx(base_dir)
    selected_template_paths: list[str] = []

    # 自我宣告 Excel（優先：資料庫&文件範本/；fallback：原目錄）
    declaration_xlsx_path = None
    for root in [assets_dir, base_dir]:
        p = root / "自我宣告.xlsx"
        if p.exists():
            declaration_xlsx_path = p
            break
    declaration_df = None
    if declaration_xlsx_path is not None:
        try:
            declaration_df = load_declarations_xlsx(str(declaration_xlsx_path))
        except Exception:
            declaration_df = None

    col_form, col_preview, col_buttons = st.columns([4, 4, 1])

    # 三欄頂部標題同一水平線（字級一致）
    with col_form:
        st.markdown("### 區塊一：基本資訊")
    with col_preview:
        st.markdown("### PDF 即時預覽")
    with col_buttons:
        st.markdown("### 檔案下載")

    chosen_date = date.today()

    with col_form:
        col_info_left, col_info_right = st.columns(2)
        with col_info_left:
            st.markdown(
                "廠商名稱 <span style='color:red'><b>* (與使用單位擇一必填)</b></span>",
                unsafe_allow_html=True,
            )
            st.text_input("廠商名稱", key="inp_vendor_name", label_visibility="collapsed")
            st.markdown("案件名稱 <span style='color:gray'>(選填)</span>", unsafe_allow_html=True)
            st.text_input("案件名稱", key="inp_project_name", label_visibility="collapsed")
            st.markdown("經銷商名稱 <span style='color:gray'>(選填)</span>", unsafe_allow_html=True)
            st.text_input("經銷商名稱", key="inp_distributor", label_visibility="collapsed")
            chosen_date = st.date_input("選擇日期 (預設今天)", value=date.today())
            d = _taiwan_date_strings(chosen_date)
            st.caption(f"西元：{d['date_greg_slash']} / 民國：{d['date_roc_slash']}")
        with col_info_right:
            st.markdown(
                "使用單位 <span style='color:red'><b>* (與廠商名稱擇一必填)</b></span>",
                unsafe_allow_html=True,
            )
            st.text_input("使用單位", key="inp_use_unit", label_visibility="collapsed")
            st.markdown("其他資訊 <span style='color:gray'>(選填)</span>", unsafe_allow_html=True)
            st.text_area("其他資訊", key="inp_other_info", height=210, label_visibility="collapsed")

        use_unit = str(st.session_state.get("inp_use_unit", "") or "")
        vendor_name = str(st.session_state.get("inp_vendor_name", "") or "")
        project_name = str(st.session_state.get("inp_project_name", "") or "")
        distributor = str(st.session_state.get("inp_distributor", "") or "")
        other_info = str(st.session_state.get("inp_other_info", "") or "")

    selected_paths_preview = st.session_state.get("selected_template_paths", []) or []
    decl_selected_preview = st.session_state.get("selected_declarations", []) or []
    options_preview = [Path(p).stem for p in selected_paths_preview]
    decl_label_preview = None
    if decl_selected_preview:
        decl_label_preview = f"{','.join(decl_selected_preview)}自我宣告"
        options_preview = [decl_label_preview] + options_preview

    with col_preview:
        if options_preview:
            current_pick_pv = st.session_state.get("pdf_preview_pick")
            default_index_pv = (
                options_preview.index(current_pick_pv) if current_pick_pv in options_preview else 0
            )
            st.selectbox(
                "選擇預覽文件",
                options=options_preview,
                index=default_index_pv,
                key="pdf_preview_pick",
            )
        else:
            st.selectbox(
                "選擇預覽文件",
                options=[""],
                index=0,
                key="pdf_preview_pick",
                disabled=True,
            )

    with col_buttons:
        st.caption("勾選文件並產生後可下載")

    with col_form:
        st.subheader("區塊二：文件類型選擇")
        st.markdown("請勾選文件類型 <span style='color:red'><b>*</b></span>", unsafe_allow_html=True)
        tabs = st.tabs(["自我宣告", "證明文件", "英文文件", "其他文件"])
        selected_declarations: list[str] = []
        for cat_name, tab in zip(["自我宣告", "證明文件", "英文文件", "其他文件"], tabs):
            with tab:
                if cat_name == "自我宣告":
                    # 依自我宣告.xlsx 的「宣告」動態產生 checkbox（每欄最多 3 列）
                    if declaration_df is None or declaration_df.empty:
                        st.caption("找不到可用的自我宣告清單（請確認有放入 自我宣告.xlsx 且包含『宣告』『中文內容』欄位）。")
                    else:
                        decl_list = declaration_df["宣告"].tolist()
                        n_decl = len(decl_list)
                        n_cols_d = max(1, math.ceil(n_decl / 3)) if n_decl else 1
                        cols_decl = st.columns(n_cols_d)
                        for i, decl in enumerate(decl_list):
                            key = f"decl_{sanitize_filename(decl)}"
                            with cols_decl[i // 3]:
                                if st.checkbox(str(decl), key=key):
                                    selected_declarations.append(str(decl))
                else:
                    cat_dict = categorized[cat_name]
                    if not cat_dict:
                        st.caption("此分類目前沒有檔案")
                    else:
                        items_tpl = list(cat_dict.items())
                        n_tpl = len(items_tpl)
                        n_cols_t = max(1, math.ceil(n_tpl / 3)) if n_tpl else 1
                        cols_tpl = st.columns(n_cols_t)
                        for i, (stem, fpath) in enumerate(items_tpl):
                            k = f"docchk_{sanitize_filename(fpath)}"
                            with cols_tpl[i // 3]:
                                if st.checkbox(stem, key=k):
                                    selected_template_paths.append(fpath)
        st.session_state.selected_template_paths = selected_template_paths
        st.session_state.selected_declarations = selected_declarations

        try:
            _miss_ctx = st.container(border=True)
        except TypeError:
            _miss_ctx = st.container()
        with _miss_ctx:
            render_missing_info_tracker(
                selected_template_paths=selected_template_paths,
                selected_declarations=selected_declarations,
                use_unit=use_unit,
                vendor_name=vendor_name,
                cart=st.session_state.cart_products,
            )

        st.subheader("區塊三：產品明細(動態購物車)")

        product_df = None
        if product_db_path and product_db_path.exists():
            try:
                product_df = load_product_db(str(product_db_path))
            except Exception as e:
                st.error(f"產品資料庫讀取失敗：{e}")
        else:
            st.error(f"找不到產品資料庫：{product_db_candidates}")

        selected_model = None
        selected_product_name = ""
        selected_unit = ""
        selected_warranty = ""

        col3_left, col3_right = st.columns([1, 2])
        with col3_left:
            opt_use_qty = st.checkbox("啟用數量與單位", value=False, key="opt_qty")
            opt_use_warranty = st.checkbox("啟用保固資訊", value=False, key="opt_warranty")

            if product_df is not None and not product_df.empty:
                models = sorted(product_df["Model"].dropna().unique().tolist())
                selected_model = st.selectbox("選擇 Model", options=models, index=0 if models else None)
                if selected_model is not None:
                    selected_row = product_df[product_df["Model"] == selected_model].head(1)
                    if not selected_row.empty:
                        r0 = selected_row.iloc[0]
                        selected_product_name = str(r0["ProductName"]).strip()
                        selected_unit = str(r0.get("單位", "") or "").strip()
                        selected_warranty = str(r0.get("保固", "") or "").strip()
            st.text_input(
                "設備名稱",
                value=(selected_product_name or "-") if selected_model is not None else "-",
                disabled=True,
            )
            if opt_use_qty:
                st.number_input("數量", min_value=1, value=1, step=1, key="add_item_qty")
                st.caption(f"單位：{selected_unit}" if selected_unit else "單位：-")
            if opt_use_warranty:
                st.caption(f"保固：{selected_warranty}" if selected_warranty else "保固：-")

            add_clicked = st.button("加入清單", use_container_width=True, key="add_to_cart_btn")

        with col3_right:
            st.markdown("#### 目前已加入的產品")
            has_cart_now = bool(st.session_state.cart_products)
            if not has_cart_now:
                st.dataframe(_empty_cart_preview_dataframe(opt_use_qty, opt_use_warranty), use_container_width=True)
            else:
                cart_df = _cart_to_editor_dataframe(st.session_state.cart_products, opt_use_qty, opt_use_warranty)
                disabled_cols = [c for c in ["項次", "Model", "設備名稱", "單位", "保固"] if c in cart_df.columns]
                edited_df = st.data_editor(
                    cart_df,
                    key="cart_products_editor",
                    use_container_width=True,
                    hide_index=True,
                    disabled=disabled_cols,
                    num_rows="fixed",
                )
                _sync_qty_from_editor(edited_df, st.session_state.cart_products, opt_use_qty)
                if st.button("清空清單", type="secondary", key="clear_cart_btn"):
                    st.session_state.cart_products = []

        # 兩欄 widget 皆已建立後再處理加入，右欄 number_input 的 session 值才與本輪一致
        if add_clicked and product_df is not None and not product_df.empty and selected_model and selected_product_name:
            exists = any(
                p.get("Model") == selected_model and p.get("ProductName") == selected_product_name
                for p in st.session_state.cart_products
            )
            if not exists:
                qty_val = 1
                if opt_use_qty:
                    try:
                        qty_val = int(st.session_state.get("add_item_qty", 1))
                    except (TypeError, ValueError):
                        qty_val = 1
                    if qty_val < 1:
                        qty_val = 1
                item = {"Model": selected_model, "ProductName": selected_product_name}
                item["數量"] = qty_val if opt_use_qty else ""
                item["單位"] = selected_unit if opt_use_qty else ""
                item["保固"] = selected_warranty if opt_use_warranty else ""
                st.session_state.cart_products.append(item)

        # --- 與右側下載共用：快取產生 docx，多份時提供 ZIP ---
        rendered_results, batch_render_error = ensure_rendered_docx_results(
            use_unit=use_unit,
            vendor_name=vendor_name,
            project_name=project_name,
            distributor=distributor,
            other_info=other_info,
            chosen_date=chosen_date,
            assets_dir=assets_dir,
            base_dir=base_dir,
            declaration_df=declaration_df,
        )
        if batch_render_error:
            st.error(f"文件渲染失敗：{batch_render_error}")

    # 中欄：PDF 即時預覽（選擇預覽文件之下拉選單已置於標題正下方，與左欄第一列對齊）
    with col_preview:
        selected_paths = st.session_state.get("selected_template_paths", []) or []
        decl_selected = st.session_state.get("selected_declarations", []) or []
        decl_label = None
        if decl_selected:
            decl_label = f"{','.join(decl_selected)}自我宣告"
        preview_pick = st.session_state.get("pdf_preview_pick") or ""

        pick_path = None
        is_decl_preview = False
        if decl_label and preview_pick == decl_label:
            # 自我宣告固定模板：規格符合證明書.docx
            for root in [assets_dir, base_dir]:
                p = root / "規格符合證明書.docx"
                if p.exists():
                    pick_path = str(p)
                    break
            is_decl_preview = True if pick_path else False
        elif preview_pick and selected_paths:
            for p in selected_paths:
                if Path(p).stem == preview_pick:
                    pick_path = p
                    break

        # 一般模板：需要購物車；自我宣告：不強制需要購物車
        has_party = bool((use_unit or "").strip() or (vendor_name or "").strip())
        has_cart = bool(st.session_state.cart_products)
        can_preview_pdf = has_party and bool(pick_path) and (has_cart or is_decl_preview)
        pdf_b64 = st.session_state.get("last_pdf_auto_b64")

        # --- PDF 即時預覽：輸入有變就自動更新 ---
        if can_preview_pdf:
            show_qty_flag = bool(st.session_state.get("opt_qty", False))
            show_warranty_flag = bool(st.session_state.get("opt_warranty", False))
            signature_payload = {
                "use_unit": use_unit,
                "vendor_name": vendor_name,
                "project_name": project_name,
                "distributor": distributor,
                "other_info": other_info,
                "decl_selected": decl_selected,
                "is_decl_preview": is_decl_preview,
                "date": str(chosen_date),
                "pick_path": str(pick_path),
                "cart": st.session_state.cart_products if has_cart else [],
                "show_qty": show_qty_flag,
                "show_warranty": show_warranty_flag,
            }
            signature = json.dumps(signature_payload, ensure_ascii=False, sort_keys=True)

            # 簡單節流：同一個 signature 在短時間內不要重複嘗試轉檔
            now_ts = time.time()
            recently_attempted = (
                st.session_state.get("last_pdf_auto_attempt_sig") == signature
                and (now_ts - float(st.session_state.get("last_pdf_auto_attempt_ts") or 0.0)) < 1.5
            )

            if (st.session_state.get("last_pdf_auto_signature") != signature) and not recently_attempted:
                st.session_state.last_pdf_auto_attempt_sig = signature
                st.session_state.last_pdf_auto_attempt_ts = now_ts
                try:
                    context = build_context(
                        use_unit=use_unit,
                        vendor_name=vendor_name,
                        project_name=project_name,
                        distributor=distributor,
                        other_info=other_info,
                        d=chosen_date,
                        products=st.session_state.cart_products if has_cart else [],
                        show_qty=show_qty_flag,
                        show_warranty=show_warranty_flag,
                    )
                    # 自我宣告（穩定方案）：
                    # - 範本放 `{{ declaration_placeholder }}`（建議放在「項目符號」段落內）
                    # - 程式用 python-docx 後處理插入多個 bullet 段落
                    if is_decl_preview and decl_selected and declaration_df is not None and (not declaration_df.empty):
                        parts = declaration_df[declaration_df["宣告"].isin(decl_selected)]["中文內容"].tolist()
                        parts = [str(x).strip() for x in parts if str(x).strip()]
                        if parts:
                            context["declaration_items"] = parts
                            context["declaration_placeholder"] = DECLARATION_BULLET_PLACEHOLDER
                    docx_bytes = render_doc(pick_path, context)
                    pdf_b64_new = docx_bytes_to_pdf_preview_base64(docx_bytes)
                    st.session_state.last_pdf_auto_b64 = pdf_b64_new
                    st.session_state.last_pdf_auto_signature = signature
                    pdf_b64 = pdf_b64_new
                except Exception:
                    st.session_state.last_pdf_auto_b64 = None
                    pdf_b64 = None
        else:
            st.session_state.last_pdf_auto_b64 = None
            st.session_state.last_pdf_auto_signature = ""
            pdf_b64 = None

        if pdf_b64:
            st.components.v1.html(
                f"""
                <div id="pdf-wrap" style="width:100%;height:760px;overflow:auto;background:#fff;"></div>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
                <script>
                (function() {{
                  const b64 = "{pdf_b64}";
                  const binary = atob(b64);
                  const len = binary.length;
                  const bytes = new Uint8Array(len);
                  for (let i = 0; i < len; i++) {{
                    bytes[i] = binary.charCodeAt(i);
                  }}
                  const blob = new Blob([bytes], {{ type: "application/pdf" }});
                  const url = URL.createObjectURL(blob);
                  const container = document.getElementById("pdf-wrap");
                  const loadingTask = window['pdfjsLib'].getDocument(url);
                  loadingTask.promise.then(function(pdf) {{
                    for (let pageNum = 1; pageNum <= pdf.numPages; pageNum++) {{
                      pdf.getPage(pageNum).then(function(page) {{
                        const viewport = page.getViewport({{ scale: 1.25 }});
                        const canvas = document.createElement('canvas');
                        const ctx = canvas.getContext('2d');
                        canvas.width = viewport.width;
                        canvas.height = viewport.height;
                        canvas.style.display = 'block';
                        canvas.style.margin = '0 auto 12px auto';
                        container.appendChild(canvas);
                        page.render({{ canvasContext: ctx, viewport: viewport }});
                      }});
                    }}
                  }});
                }})();
                </script>
                """,
                height=780,
            )
        else:
            st.markdown(
                "<div style='height:760px;background:#fff;border:1px dashed #e5e7eb;border-radius:10px;'></div>",
                unsafe_allow_html=True,
            )

    # 右欄：依「選擇預覽文件」下載 Word / PDF / PNG；最下方為多份 ZIP
    with col_buttons:
        sel_paths_btn = st.session_state.get("selected_template_paths", []) or []
        decl_sel_btn = st.session_state.get("selected_declarations", []) or []
        decl_lbl_btn = f"{','.join(decl_sel_btn)}自我宣告" if decl_sel_btn else None
        pick_btn = st.session_state.get("pdf_preview_pick") or ""
        r_key_btn = result_key_for_preview_pick(pick_btn, decl_lbl_btn, decl_sel_btn, sel_paths_btn)
        payload_btn = rendered_results.get(r_key_btn) if r_key_btn else None

        if not payload_btn:
            st.caption("請在左側勾選要產生的文件；一般模板需加入購物車後才可下載。")
        else:
            docx_b = payload_btn["bytes"]
            base_stem = Path(payload_btn["filename"]).stem
            st.download_button(
                label="下載 Word",
                data=docx_b,
                file_name=f"{base_stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="col_dl_docx",
            )
            pdf_b, pdf_err = docx_bytes_to_pdf_bytes(docx_b)
            if pdf_b:
                st.download_button(
                    label="下載 PDF",
                    data=pdf_b,
                    file_name=f"{base_stem}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key="col_dl_pdf",
                )
            else:
                st.caption(pdf_err or "無法產生 PDF")
            if pdf_b:
                png_b, png_err = pdf_bytes_to_png_bytes(pdf_b)
            else:
                png_b, png_err = None, None
            if png_b:
                st.download_button(
                    label="下載 PNG",
                    data=png_b,
                    file_name=f"{base_stem}.png",
                    mime="image/png",
                    use_container_width=True,
                    key="col_dl_png",
                )
            elif pdf_b:
                st.caption(png_err or "無法產生 PNG")

        st.divider()
        if len(rendered_results) > 1:
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                for _k, payload in rendered_results.items():
                    zf.writestr(payload["filename"], payload["bytes"])
            st.download_button(
                label="下載多份文件壓縮檔 (word)",
                data=zip_buffer.getvalue(),
                file_name="ACTi_文件包.zip",
                mime="application/zip",
                use_container_width=True,
                key="col_dl_zip_bundle",
            )


if __name__ == "__main__":
    main()

